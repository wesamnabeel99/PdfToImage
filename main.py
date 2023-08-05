from docx import Document
from docx.shared import Pt
from docx2pdf import convert
import os
from processing.pdf_processor.pdf_crop import crop_pdf,extract_images_from_pdf
import Augmentor
def write_in_document(doc,text,font_name):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.size=Pt(16)
    run.font.name= font_name
    run.bidi = 'RTL'


arabic_sentences = [
    "أهلاً بك في عالم الإبداع والتحديات الجديدة.",
    "بين أوراق الكتب نجد الكنوز الثمينة للمعرفة.",
    "تعلم وتطور، فالمعرفة لا تعرف حدودًا.",
    "جرب وفشل، فالنجاح يأتي بعد العزيمة.",
    "خلف كل صعوبة يكمن فرصة جديدة للتعلم.",
    "رحلة الألف ميل تبدأ بخطوة واحدة.",
    "زرع المحبة والتسامح يجني ثمار السلام.",
    "سواء كنت تعتقد أنك تستطيع أو لا تستطيع، فأنت على حق.",
    "شارك المعرفة ولا تبخل بها على الآخرين.",
    "صبرًا وعزيمة، تتحقق الأحلام والطموحات.",
    "ضع خططًا وأهدافًا لتحقيق التقدم والنجاح.",
    "طريق النجاح مليء بالتحديات والفرص المتجددة.",
    "ظل الصبر طريقًا للفوز وتحقيق الإنجازات.",
    "عليك أن تصنع فرصك وتجعل من الصعب سهلًا.",
    "في قلب الصحراء، تكمن جمال الطبيعة الخلابة.",
    "قم بما يجب عليك القيام به، ودع النتائج تتولد بشكل طبيعي.",
    "كل شيء ممكن إذا كنت مصممًا وملتزمًا بتحقيقه.",
    "لن يُكافئ النجاح إلا الذين يتحملون الفشل.",
    "مع التعلم المستمر، تستمر رحلة النمو الشخصي.",
    "نحو عالم مليء بالتفاؤل والإيجابية، نسعى دومًا.",
]

arabic_letters = ["أ", "ب", "ت", "ث", "ج", "ح", "خ", "د", "ذ", "ر", "ز", "س", "ش", "ص", "ض", "ط", "ظ", "ع", "غ", "ف", "ق", "ك", "ل", "م", "ن", "هـ", "و", "ي"]
fonts = [
    "Urdu Typesetting",
    "Aldhabi",
    "B Arabic Style",
    "THARWATEMARARUQAA",
    "K Kamran",
    "A Hemmat",
    "B Shekari",
    "Dast Nevis",
    "FS_Hand_Style",
    "Ghalam1",
    "Hesham Free"
]

for font in fonts:
    doc = Document()
    for i, letter in enumerate(arabic_sentences):
        write_in_document(doc,letter,font)
        if i != len(arabic_sentences) - 1:
            doc.add_page_break()

    doc.save(f"{font}.docx")
    del doc
    convert(f"{font}.docx",f"{font}.pdf")
    crop_pdf(f"{font}.pdf",f"{font}_output.pdf")
    extract_images_from_pdf(f"{font}_output.pdf",f"output_images/{font}")


    os.remove(f"{font}.docx")
    os.remove(f"{font}.pdf")
    os.remove(f"{font}_output.pdf")

    pipeline = Augmentor.Pipeline(f"output_images/{font}")

    pipeline.rotate_without_crop(probability=0.7, max_left_rotation=10, max_right_rotation=10)
    pipeline.gaussian_distortion(probability=0.5, grid_width=7, grid_height=7, magnitude=8, corner='bell', method='in')
    pipeline.random_distortion(probability=0.5, grid_width=7, grid_height=7, magnitude=8)

    pipeline.sample(100)